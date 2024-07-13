from openai import OpenAI
from typing import Generator, List
import time
import io
from docx import Document
from docx.shared import Inches
from tinychat.utils.secrets import get_secret
from tinychat.settings import OPENAI_API_KEY_NAME
import json


class OpenAIHandler:
    def __init__(self, temperature: float = 0.0):
        self.client = OpenAI(api_key=get_secret(OPENAI_API_KEY_NAME))
        self.temperature = temperature

    def analyze_documents(self, system_prompt: str, user_input: str) -> List[dict]:
        response = self.client.chat.completions.create(
            model="gpt-4-turbo-preview",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_input},
            ],
            temperature=self.temperature,
        )
        content = response.choices[0].message.content
        print("Raw OpenAI response:")
        print(content)
        try:
            # Extract the JSON part of the response
            json_start = content.find("[")
            json_end = content.rfind("]") + 1
            if json_start != -1 and json_end != -1:
                json_content = content[json_start:json_end]
                return json.loads(json_content)
            else:
                raise ValueError("No valid JSON found in the response")
        except json.JSONDecodeError:
            raise ValueError("The AI response was not in the expected JSON format.")

    def revise_docx(self, original_path: str, revised_text: str) -> str:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Load the original document
        original_doc = Document(original_path)

        # Split the revised text into paragraphs
        revised_paragraphs = revised_text.split("\n")

        # Create a new comparison document
        comparison_doc = Document()

        # Add a title
        title = comparison_doc.add_paragraph("NDA Comparison")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].bold = True
        title.runs[0].font.size = Pt(16)

        # Add a legend
        legend = comparison_doc.add_paragraph()
        legend.add_run("Original Text").font.color.rgb = RGBColor(255, 0, 0)
        legend.add_run(" | ")
        legend.add_run("Revised Text").font.color.rgb = RGBColor(0, 0, 255)

        # Iterate through the paragraphs
        for i, original_para in enumerate(original_doc.paragraphs):
            comparison_para = comparison_doc.add_paragraph()

            # Add original text
            original_run = comparison_para.add_run(original_para.text)
            original_run.font.color.rgb = RGBColor(255, 0, 0)  # Red for original

            comparison_para.add_run(" | ")

            # Add revised text
            if i < len(revised_paragraphs):
                revised_run = comparison_para.add_run(revised_paragraphs[i])
                revised_run.font.color.rgb = RGBColor(0, 0, 255)  # Blue for revised
            else:
                revised_run = comparison_para.add_run("[No changes]")

            # Add a line break between paragraphs
            comparison_doc.add_paragraph()

        # Save the comparison document
        comparison_path = original_path.replace(".docx", "_comparison.docx")
        comparison_doc.save(comparison_path)

        return comparison_path

    def stream_response(self, user_input: str) -> Generator[str, None, None]:
        self.client.beta.threads.messages.create(
            thread_id=self.thread.id, role="user", content=user_input
        )

        run = self.client.beta.threads.runs.create(
            thread_id=self.thread.id, assistant_id=self.assistant_id
        )

        while run.status not in ["completed", "failed", "cancelled"]:
            time.sleep(0.5)
            run = self.client.beta.threads.runs.retrieve(
                thread_id=self.thread.id, run_id=run.id
            )

        if run.status != "completed":
            yield f"An error occurred while processing your request. Status: {run.status}"
            return

        messages = self.client.beta.threads.messages.list(
            thread_id=self.thread.id, order="desc", limit=1
        )

        for message in messages:
            if message.role == "assistant":
                for content in message.content:
                    if content.type == "text":
                        yield content.text.value

    def upload_file(self, file_path: str) -> str:
        with open(file_path, "rb") as file:
            response = self.client.files.create(file=file, purpose="assistants")
        return response.id

    def create_vector_store(self, name: str, file_ids: List[str]):
        self.vector_store = self.client.beta.vector_stores.create(
            name=name, file_ids=file_ids
        )
        return self.vector_store

    def update_assistant_with_vector_store(self):
        if not self.vector_store:
            raise ValueError(
                "Vector store not created. Call create_vector_store first."
            )
        self.client.beta.assistants.update(
            assistant_id=self.assistant_id,
            tool_resources={
                "file_search": {"vector_store_ids": [self.vector_store.id]}
            },
        )

    def export_conversation(self) -> str:
        messages = self.client.beta.threads.messages.list(thread_id=self.thread.id)
        conversation = ""
        for message in reversed(list(messages)):
            role = "You" if message.role == "user" else "Assistant"
            content = message.content[0].text.value if message.content else ""
            conversation += f"{role}: {content}\n\n"
        return conversation
