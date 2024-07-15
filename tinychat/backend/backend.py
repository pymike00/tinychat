from typing import Optional, List
import io
from docx import Document
import os
from tkinter import filedialog
import json
import difflib

from tinychat.llms.base import LLMProtocol
from tinychat.llms.openai import OpenAIHandler
from tinychat.utils.secrets import get_secret, set_secret


class Backend:
    def __init__(self) -> None:
        self.temperature: float = self.get_default_temperature()
        self._models = {
            "Language Model ": lambda: None,
            "OpenAI GPT-4": lambda: OpenAIHandler(self.temperature),
        }
        self._llm: Optional[LLMProtocol] = None
        self.nda_content = None
        self.guidelines = None
        self.revised_nda = None
        self.suggested_changes = None

    def available_models(self) -> list:
        return list(self._models.keys())

    def set_model(self, model_name: str) -> None:
        if model_name not in self.available_models():
            raise KeyError(f"Invalid Model Name {model_name}")
        try:
            self._llm = self._models[model_name]()
        except ValueError as e:
            raise ValueError(
                f"Initialization Error. Have you set the API Key for {model_name}? {e}"
            )

    def get_default_temperature(self):
        """
        Get the value of temperature from tinychat.json if available
        else return default_temperature.
        """      
        default_temperature = 0.7
        try:
            temperature = float(get_secret("temperature"))
        except ValueError:
            temperature = default_temperature
            set_secret("temperature", default_temperature)
        return temperature

    def get_stream_response(self, user_input: str):
        if self._llm is None:
            raise ValueError("No Language Model Has Been Selected.")
        return self._llm.stream_response(user_input)

    def export_conversation(self):
        if self._llm is None:
            return
        new_file = filedialog.asksaveasfilename(
            initialfile="Untitled.txt",
            defaultextension=".txt",
            filetypes=[("Text File", "*.txt")],
        )
        if not new_file:
            return
        with open(new_file, "w") as f:
            f.write(self._llm.export_conversation())

    def upload_nda(self):
        file_path = filedialog.askopenfilename(
            filetypes=[("Word Document", "*.docx"), ("Text File", "*.txt")]
        )
        if not file_path:
            return "No file selected"
        
        with open(file_path, 'rb') as file:
            if file_path.endswith('.docx'):
                doc = Document(file)
                content = "\n".join([para.text for para in doc.paragraphs])
            else:
                content = file.read().decode('utf-8')
        
        self.nda_content = content
        return f"NDA uploaded successfully: {os.path.basename(file_path)}"

    def upload_guidelines(self):
        file_path = filedialog.askopenfilename(
            filetypes=[("Text File", "*.txt"), ("Word Document", "*.docx")]
        )
        if not file_path:
            return "No file selected"
        
        with open(file_path, 'rb') as file:
            if file_path.endswith('.docx'):
                doc = Document(file)
                content = "\n".join([para.text for para in doc.paragraphs])
            else:
                content = file.read().decode('utf-8')
        
        self.guidelines = content
        return f"Guidelines uploaded successfully: {os.path.basename(file_path)}"

    def download_revised_nda(self):
        if not hasattr(self, 'revised_nda'):
            return "No revised NDA available. Please analyze and revise the NDA first."
        
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")]
        )
        if not file_path:
            return "Download cancelled"
        
        doc = Document()
        for paragraph in self.revised_nda.split('\n'):
            doc.add_paragraph(paragraph)
        doc.save(file_path)
        
        return f"Revised NDA downloaded successfully: {os.path.basename(file_path)}"

    def analyze_and_revise_nda(self):
        if not hasattr(self, 'nda_content') or not hasattr(self, 'guidelines'):
            return "Please upload both NDA and guidelines before analyzing."
        
        if self._llm is None:
            raise ValueError("No Language Model Has Been Selected.")
        
        system_prompt = """You are an AI assistant specialized in analyzing and revising Non-Disclosure Agreements (NDAs). 
        Your task is to review the provided NDA and suggest revisions based on the given guidelines. 
        Provide your analysis and revisions in a clear, structured format."""

        user_input = f"""Please analyze and revise the following NDA according to these guidelines:

        Guidelines:
        {self.guidelines}

        NDA:
        {self.nda_content}

        Provide your analysis and revised NDA as a list of JSON objects, where each object represents a suggested change:
        [
            {{
                "original_text": "Full original paragraph from the NDA",
                "suggested_change": "Full revised paragraph with suggested changes",
                "justification": "Explanation for why this change was suggested"
            }},
            // ... more objects for other paragraphs
        ]
        Only include paragraphs that need changes.
        """

        try:
            response = self._llm.analyze_documents(system_prompt, user_input)
            
            if isinstance(response, list) and len(response) > 0:
                if "raw_content" in response[0]:
                    return f"Analysis complete, but the response was not in the expected format. Here's the raw response:\n\n{response[0]['raw_content']}"
                else:
                    self.suggested_changes = response
                    return "Analysis complete. Ready to review changes."
            else:
                raise ValueError("Unexpected response format from the language model.")
        except Exception as e:
            return f"An error occurred during analysis: {str(e)}"

    def review_changes(self):
        if not hasattr(self, 'suggested_changes'):
            return "No changes to review. Please analyze the NDA first."
        
        for change in self.suggested_changes:
            yield change

    def apply_approved_changes(self, approved_changes):
        if not hasattr(self, 'nda_content'):
            return "No NDA content found. Please upload an NDA first."
        
        revised_nda = self.nda_content
        for change in approved_changes:
            revised_nda = revised_nda.replace(change['original_text'], change['suggested_change'])
        
        self.revised_nda = revised_nda
        return "Changes applied successfully. You can now download the revised NDA."

    def set_system_prompt(self, prompt: str):
        if isinstance(self._llm, OpenAIHandler):
            self._llm.set_system_prompt(prompt)

    def send_message(self, user_input: str) -> str:
        if self._llm is None:
            raise ValueError("No Language Model Has Been Selected.")
        return self._llm.get_response(user_input)